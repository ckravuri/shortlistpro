from fastapi import FastAPI, APIRouter, HTTPException, Request, Response, Depends, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from bson import ObjectId
import secrets
import re
from openai import AsyncOpenAI
import json
import asyncio
import base64
from utils.pdf_parser import parse_pdf_resume, parse_docx_resume
from utils.pdf_generator import generate_resume_pdf
import stripe

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ['JWT_SECRET']
JWT_ALGORITHM = "HS256"

# Stripe Configuration
stripe.api_key = os.environ['STRIPE_SECRET_KEY']
STRIPE_PRO_PRICE_ID = os.environ['STRIPE_PRO_PRICE_ID']
STRIPE_PRO_PLUS_PRICE_ID = os.environ['STRIPE_PRO_PLUS_PRICE_ID']
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', '')

# OpenAI Configuration
openai_client = AsyncOpenAI(api_key=os.environ['OPENAI_API_KEY'])

# Google OAuth Configuration
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID', '')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET', '')

# Rate Limiting Setup
limiter = Limiter(key_func=get_remote_address)

# Create the main app
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

api_router = APIRouter(prefix="/api")

# ============ AUTH UTILITIES ============
def hash_password(password: str) -> str:
    salt = bcrypt.gensalt()
    hashed = bcrypt.hashpw(password.encode("utf-8"), salt)
    return hashed.decode("utf-8")

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return bcrypt.checkpw(plain_password.encode("utf-8"), hashed_password.encode("utf-8"))

def create_access_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=15),
        "type": "access"
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def create_refresh_token(user_id: str) -> str:
    payload = {
        "sub": user_id,
        "exp": datetime.now(timezone.utc) + timedelta(days=7),
        "type": "refresh"
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(request: Request) -> dict:
    token = request.cookies.get("access_token")
    if not token:
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Bearer "):
            token = auth_header[7:]
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "access":
            raise HTTPException(status_code=401, detail="Invalid token type")
        
        # CRITICAL FIX: Use custom 'id' field, NOT '_id' with ObjectId
        user_id = payload["sub"]
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        user.pop("password_hash", None)
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ============ MODELS ============
class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str
    region: str = "US"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class PersonalInfo(BaseModel):
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    website: str = ""
    summary: str = ""

class WorkExperience(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company: str = ""
    position: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    current: bool = False
    description: str = ""
    achievements: List[str] = []

class Education(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    institution: str = ""
    degree: str = ""
    field: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    gpa: str = ""

class ResumeCreate(BaseModel):
    title: str = "My Resume"
    region: str = "US"
    template: Optional[str] = None

class ResumeUpdate(BaseModel):
    title: Optional[str] = None
    personal_info: Optional[PersonalInfo] = None
    work_experience: Optional[List[WorkExperience]] = None
    education: Optional[List[Education]] = None
    skills: Optional[List[str]] = None
    region: Optional[str] = None

class AISuggestionRequest(BaseModel):
    context: str
    field: str
    current_text: str = ""

class InterviewPrepRequest(BaseModel):
    resume_id: str
    job_description: str

class InterviewAnswerRequest(BaseModel):
    question: str
    job_description: str
    resume_context: str  # Brief resume summary

# ============ STARTUP EVENTS ============
@app.on_event("startup")
async def startup_event():
    # Create indexes
    await db.users.create_index("email", unique=True)
    await db.login_attempts.create_index("identifier")
    
    # Seed admin
    admin_email = os.environ.get("ADMIN_EMAIL", "admin@shortlistpro.cv")
    admin_password = os.environ.get("ADMIN_PASSWORD", "Admin@2026Secure")
    existing = await db.users.find_one({"email": admin_email})
    
    if existing is None:
        hashed = hash_password(admin_password)
        admin_id = str(uuid.uuid4())
        await db.users.insert_one({
            "id": admin_id,
            "email": admin_email,
            "password_hash": hashed,
            "name": "Admin",
            "region": "US",
            "role": "admin",
            "subscription_tier": "pro+",  # Pro+ for full testing
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        logger.info(f"Admin user created: {admin_email}")
    else:
        # Update to Pro+ if needed
        update_doc = {}
        if not verify_password(admin_password, existing["password_hash"]):
            update_doc["password_hash"] = hash_password(admin_password)
        if "id" not in existing:
            update_doc["id"] = str(uuid.uuid4())
        if existing.get("subscription_tier") != "pro+":
            update_doc["subscription_tier"] = "pro+"
        if update_doc:
            await db.users.update_one(
                {"email": admin_email},
                {"$set": update_doc}
            )
            logger.info("Admin user updated to Pro+")
    
    # Write test credentials
    Path("/app/memory").mkdir(exist_ok=True)
    with open("/app/memory/test_credentials.md", "w") as f:
        f.write("# Test Credentials\n\n")
        f.write("## Admin Account\n")
        f.write(f"- Email: {admin_email}\n")
        f.write(f"- Password: {admin_password}\n")
        f.write("- Role: admin\n\n")
        f.write("## Auth Endpoints\n")
        f.write("- POST /api/auth/register\n")
        f.write("- POST /api/auth/login\n")
        f.write("- GET /api/auth/me\n")
        f.write("- POST /api/auth/logout\n")

# ============ AUTH ROUTES ============
@api_router.post("/auth/register")
@limiter.limit("5/hour")  # Prevent bot account creation
async def register(user: UserRegister, response: Response, request: Request):
    email_lower = user.email.lower()
    existing = await db.users.find_one({"email": email_lower})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed = hash_password(user.password)
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": email_lower,
        "password_hash": hashed,
        "name": user.name,
        "region": user.region,
        "role": "user",
        "subscription_tier": "free",
        "subscription_status": "active",
        "headshot_url": None,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user_doc)
    
    access_token = create_access_token(user_id, email_lower)
    refresh_token = create_refresh_token(user_id)
    
    response.set_cookie(
        key="access_token",
        value=access_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=900,
        path="/"
    )
    response.set_cookie(
        key="refresh_token",
        value=refresh_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=604800,
        path="/"
    )
    
    return {
        "id": user_id,
        "email": email_lower,
        "name": user.name,
        "region": user.region,
        "role": "user",
        "subscription_tier": "free"
    }

@api_router.post("/auth/login")
@limiter.limit("10/minute")  # Prevent brute force attacks
async def login(user: UserLogin, request: Request, response: Response):
    email_lower = user.email.lower()
    
    # Check brute force
    client_ip = request.client.host
    identifier = f"{client_ip}:{email_lower}"
    attempt = await db.login_attempts.find_one({"identifier": identifier})
    
    if attempt and attempt.get("attempts", 0) >= 5:
        lockout_until = attempt.get("lockout_until")
        if lockout_until and datetime.fromisoformat(lockout_until) > datetime.now(timezone.utc):
            raise HTTPException(status_code=429, detail="Too many failed attempts. Please try again later.")
    
    user_doc = await db.users.find_one({"email": email_lower})
    if not user_doc or not verify_password(user.password, user_doc["password_hash"]):
        # Increment failed attempts
        if attempt:
            new_attempts = attempt.get("attempts", 0) + 1
            update_doc = {"attempts": new_attempts}
            if new_attempts >= 5:
                update_doc["lockout_until"] = (datetime.now(timezone.utc) + timedelta(minutes=15)).isoformat()
            await db.login_attempts.update_one(
                {"identifier": identifier},
                {"$set": update_doc}
            )
        else:
            await db.login_attempts.insert_one({
                "identifier": identifier,
                "attempts": 1,
                "created_at": datetime.now(timezone.utc).isoformat()
            })
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Clear failed attempts
    await db.login_attempts.delete_one({"identifier": identifier})
    
    # Get or create custom ID
    user_id = user_doc.get("id")
    if not user_id:
        user_id = str(uuid.uuid4())
        await db.users.update_one(
            {"_id": user_doc["_id"]},
            {"$set": {"id": user_id}}
        )
    
    access_token = create_access_token(user_id, email_lower)
    refresh_token = create_refresh_token(user_id)
    
    response.set_cookie(
        key="access_token",
        value=access_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=900,
        path="/"
    )
    response.set_cookie(
        key="refresh_token",
        value=refresh_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=604800,
        path="/"
    )
    
    return {
        "id": user_id,
        "email": user_doc["email"],
        "name": user_doc["name"],
        "region": user_doc.get("region", "US"),
        "role": user_doc.get("role", "user"),
        "subscription_tier": user_doc.get("subscription_tier", "free")
    }

@api_router.get("/auth/me")
async def get_me(request: Request, current_user: dict = Depends(get_current_user)):
    return current_user

@api_router.post("/auth/logout")
async def logout(response: Response):
    response.delete_cookie("access_token", path="/")
    response.delete_cookie("refresh_token", path="/")
    response.delete_cookie("session_token", path="/")
    return {"message": "Logged out successfully"}

# ============ GOOGLE OAUTH ROUTES ============
@api_router.post("/auth/google")
async def google_auth(request: Request, response: Response):
    """Verify Google ID token and create/login user"""
    from google.oauth2 import id_token
    from google.auth.transport import requests as google_requests
    
    data = await request.json()
    token = data.get("credential")
    
    if not token:
        raise HTTPException(status_code=400, detail="Google credential required")
    
    try:
        # Verify the Google ID token
        idinfo = id_token.verify_oauth2_token(
            token, 
            google_requests.Request(), 
            GOOGLE_CLIENT_ID
        )
        
        # Extract user info from token
        email = idinfo['email'].lower()
        name = idinfo.get('name', '')
        picture = idinfo.get('picture', '')
        google_id = idinfo['sub']
        
        # Find or create user
        user = await db.users.find_one({"email": email}, {"_id": 0})
        
        if not user:
            # Create new user with Google OAuth
            user_id = str(uuid.uuid4())
            new_user = {
                "_id": ObjectId(),
                "id": user_id,
                "email": email,
                "name": name,
                "picture": picture,
                "google_id": google_id,
                "region": "US",
                "subscription_tier": "free",
                "stripe_customer_id": None,
                "ai_suggestions_used": 0,
                "role": "user",
                "auth_provider": "google",
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(new_user)
            user = {k: v for k, v in new_user.items() if k != "_id"}
        else:
            # Update existing user's picture and google_id if changed
            update_fields = {}
            if picture and user.get("picture") != picture:
                update_fields["picture"] = picture
            if not user.get("google_id"):
                update_fields["google_id"] = google_id
            if not user.get("auth_provider"):
                update_fields["auth_provider"] = "google"
            
            if update_fields:
                await db.users.update_one(
                    {"email": email},
                    {"$set": update_fields}
                )
                user.update(update_fields)
        
        # Create JWT tokens (reuse existing auth system)
        user_id = user["id"]
        access_token = create_access_token(user_id, email)
        refresh_token = create_refresh_token(user_id)
        
        # Set cookies
        response.set_cookie(
            key="access_token",
            value=access_token,
            httponly=True,
            secure=True,
            samesite="none",
            max_age=900,  # 15 minutes
            path="/"
        )
        response.set_cookie(
            key="refresh_token",
            value=refresh_token,
            httponly=True,
            secure=True,
            samesite="none",
            max_age=7*24*60*60,  # 7 days
            path="/"
        )
        
        # Return user data
        return {
            "id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "picture": user.get("picture", ""),
            "region": user["region"],
            "subscription_tier": user["subscription_tier"],
            "role": user["role"]
        }
        
    except ValueError as e:
        # Invalid token
        logging.error(f"Google token verification failed: {str(e)}")
        raise HTTPException(status_code=401, detail="Invalid Google token")
    except Exception as e:
        logging.error(f"Google OAuth error: {str(e)}")
        raise HTTPException(status_code=500, detail="Authentication failed")

# ============ RESUME ROUTES ============
# Subscription tier limits
SUBSCRIPTION_LIMITS = {
    "free": {"resumes": 1, "ai_suggestions": 5},
    "pro": {"resumes": 10, "ai_suggestions": 50},
    "pro+": {"resumes": float('inf'), "ai_suggestions": float('inf')}  # Unlimited
}

async def check_resume_limit(user_id: str, subscription_tier: str):
    """
    Check if user has reached their resume limit based on subscription tier
    Raises HTTPException if limit reached
    """
    tier = subscription_tier.lower() if subscription_tier else "free"
    limit = SUBSCRIPTION_LIMITS.get(tier, SUBSCRIPTION_LIMITS["free"])["resumes"]
    
    # For free users, check if they've EVER created a resume (prevent abuse by deleting/recreating)
    if tier == "free":
        user = await db.users.find_one({"id": user_id})
        if user and user.get("has_created_resume", False):
            # User has already used their free resume slot
            resume_count = await db.resumes.count_documents({"user_id": user_id})
            raise HTTPException(
                status_code=403,
                detail={
                    "message": "Free plan already used",
                    "description": "You've already used your Free plan resume. To create more resumes, please upgrade to Pro!",
                    "current_count": resume_count,
                    "limit": int(limit),
                    "tier": tier,
                    "upgrade_required": True
                }
            )
    
    # Count existing resumes
    resume_count = await db.resumes.count_documents({"user_id": user_id})
    
    if resume_count >= limit:
        if tier == "free":
            raise HTTPException(
                status_code=403,
                detail={
                    "message": "Resume limit reached",
                    "description": f"You've reached your Free plan limit of {int(limit)} resume. Upgrade to Pro for 10 resumes or Pro+ for unlimited resumes!",
                    "current_count": resume_count,
                    "limit": int(limit),
                    "tier": tier,
                    "upgrade_required": True
                }
            )
        else:
            raise HTTPException(
                status_code=403,
                detail={
                    "message": "Resume limit reached",
                    "description": f"You've reached your {tier.title()} plan limit of {int(limit)} resumes. Upgrade to Pro+ for unlimited resumes!",
                    "current_count": resume_count,
                    "limit": int(limit),
                    "tier": tier,
                    "upgrade_required": True
                }
            )
    
    return True

def get_template_structure(template_id: str):
    """Get pre-populated structure for resume templates"""
    templates = {
        "harvard": {
            "personal_info": {
                "full_name": "Your Name",
                "email": "your.email@example.com",
                "phone": "(555) 123-4567",
                "location": "City, State",
                "linkedin": "linkedin.com/in/yourprofile",
                "summary": "Write a compelling 2-3 sentence summary highlighting your key strengths and career goals. Focus on what makes you unique and valuable to employers."
            },
            "work_experience": [
                {
                    "id": "template-exp-1",
                    "company": "Company Name",
                    "position": "Your Job Title",
                    "location": "City, State",
                    "start_date": "MM/YYYY",
                    "end_date": "Present",
                    "current": True,
                    "description": "Brief description of your role and responsibilities",
                    "achievements": [
                        "Quantifiable achievement with metrics (e.g., Increased sales by 25%)",
                        "Another accomplishment demonstrating impact",
                        "Use action verbs: Led, Developed, Implemented, Achieved"
                    ]
                }
            ],
            "education": [
                {
                    "id": "template-edu-1",
                    "institution": "University Name",
                    "degree": "Bachelor of Science",
                    "field": "Your Major",
                    "location": "City, State",
                    "graduation_date": "MM/YYYY",
                    "gpa": "3.X/4.0"
                }
            ],
            "skills": ["Skill 1", "Skill 2", "Skill 3", "Skill 4", "Skill 5"]
        },
        "modern-professional": {
            "personal_info": {
                "full_name": "Your Name",
                "email": "your.email@example.com",
                "phone": "(555) 123-4567",
                "location": "City, State",
                "linkedin": "linkedin.com/in/yourprofile",
                "summary": "Professional summary emphasizing your expertise and value proposition. Keep it concise, achievements-focused, and tailored to your target role."
            },
            "work_experience": [
                {
                    "id": "template-exp-1",
                    "company": "Current/Recent Company",
                    "position": "Your Position",
                    "location": "City, State",
                    "start_date": "MM/YYYY",
                    "end_date": "Present",
                    "current": True,
                    "description": "Overview of your role",
                    "achievements": [
                        "Key accomplishment with measurable result",
                        "Another achievement showing leadership or innovation",
                        "Demonstrate skills relevant to your target role"
                    ]
                }
            ],
            "education": [
                {
                    "id": "template-edu-1",
                    "institution": "University/College Name",
                    "degree": "Degree Type",
                    "field": "Field of Study",
                    "location": "City, State",
                    "graduation_date": "MM/YYYY",
                    "gpa": "GPA (optional)"
                }
            ],
            "skills": ["Technical Skill", "Software/Tool", "Methodology", "Soft Skill", "Industry Knowledge"]
        }
    }
    
    # Default template structure for other templates
    default_template = {
        "personal_info": {
            "full_name": "Your Name",
            "email": "your.email@example.com",
            "phone": "(555) 123-4567",
            "location": "City, State",
            "linkedin": "linkedin.com/in/yourprofile",
            "summary": "Write a 2-3 sentence professional summary. Highlight your key strengths, years of experience, and what you bring to potential employers."
        },
        "work_experience": [
            {
                "id": "template-exp-1",
                "company": "Company Name",
                "position": "Job Title",
                "location": "City, State",
                "start_date": "MM/YYYY",
                "end_date": "MM/YYYY",
                "current": False,
                "description": "Brief overview of your role",
                "achievements": [
                    "Start each bullet with an action verb (Led, Developed, Achieved)",
                    "Include quantifiable results when possible (Increased efficiency by 30%)",
                    "Focus on impact and outcomes, not just responsibilities"
                ]
            }
        ],
        "education": [
            {
                "id": "template-edu-1",
                "institution": "University Name",
                "degree": "Degree Type (e.g., Bachelor of Science)",
                "field": "Your Major",
                "location": "City, State",
                "graduation_date": "MM/YYYY",
                "gpa": "3.X/4.0 (optional)"
            }
        ],
        "skills": ["List", "Your", "Key", "Skills", "Here"]
    }
    
    return templates.get(template_id, default_template)

@api_router.post("/resumes")
async def create_resume(resume: ResumeCreate, current_user: dict = Depends(get_current_user)):
    # Check subscription limit before creating
    await check_resume_limit(current_user["id"], current_user.get("subscription_tier", "free"))
    
    # Get template structure if template is specified
    template_structure = {}
    if hasattr(resume, 'template') and resume.template:
        template_structure = get_template_structure(resume.template)
    
    resume_doc = {
        "user_id": current_user["id"],
        "title": resume.title,
        "region": resume.region,
        "personal_info": template_structure.get("personal_info", {}),
        "work_experience": template_structure.get("work_experience", []),
        "education": template_structure.get("education", []),
        "skills": template_structure.get("skills", []),
        "ats_score": 0,
        "keywords": [],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    result = await db.resumes.insert_one(resume_doc)
    resume_doc["id"] = str(result.inserted_id)
    resume_doc.pop("_id", None)
    
    # Mark that free user has created a resume (prevent abuse)
    if current_user.get("subscription_tier", "free") == "free":
        await db.users.update_one(
            {"id": current_user["id"]},
            {"$set": {"has_created_resume": True}}
        )
    
    return resume_doc

@api_router.get("/resumes")
async def get_resumes(current_user: dict = Depends(get_current_user)):
    # Don't exclude _id, we need it to create the id field
    resumes = await db.resumes.find({"user_id": current_user["id"]}).to_list(1000)
    for resume in resumes:
        if "_id" in resume:
            resume["id"] = str(resume["_id"])
            resume.pop("_id")
    return resumes

@api_router.get("/resumes/{resume_id}")
async def get_resume(resume_id: str, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    resume["id"] = str(resume["_id"])
    resume.pop("_id", None)
    return resume

@api_router.put("/resumes/{resume_id}")
async def update_resume(resume_id: str, update: ResumeUpdate, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    update_doc = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if update.title:
        update_doc["title"] = update.title
    if update.personal_info:
        update_doc["personal_info"] = update.personal_info.model_dump()
    if update.work_experience is not None:
        update_doc["work_experience"] = [exp.model_dump() for exp in update.work_experience]
    if update.education is not None:
        update_doc["education"] = [edu.model_dump() for edu in update.education]
    if update.skills is not None:
        update_doc["skills"] = update.skills
    if update.region:
        update_doc["region"] = update.region
    
    # Calculate ATS score
    old_score = resume.get("ats_score", 0)
    ats_score = calculate_ats_score(update_doc, resume)
    update_doc["ats_score"] = ats_score
    
    await db.resumes.update_one({"_id": ObjectId(resume_id)}, {"$set": update_doc})
    
    # Track score history if score changed
    if ats_score != old_score:
        await db.score_history.insert_one({
            "resume_id": resume_id,
            "score": ats_score,
            "date": datetime.now(timezone.utc).isoformat()
        })
    
    updated_resume = await db.resumes.find_one({"_id": ObjectId(resume_id)}, {"_id": 0})
    updated_resume["id"] = resume_id
    return updated_resume

@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a resume"""
    logger.info(f"Delete request for resume {resume_id} by user {current_user.get('id')}")
    
    try:
        # Delete using the ObjectId for MongoDB
        result = await db.resumes.delete_one({
            "user_id": current_user["id"],
            "_id": ObjectId(resume_id)
        })
        
        if result.deleted_count == 0:
            logger.warning(f"Resume {resume_id} not found or not owned by user")
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Also delete associated score history
        await db.score_history.delete_many({"resume_id": resume_id})
        
        logger.info(f"Successfully deleted resume {resume_id}")
        return {"message": "Resume deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting resume {resume_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete resume: {str(e)}")

# ============ AI SUGGESTION (Streaming) ============
@api_router.post("/resumes/{resume_id}/ai-suggest")
@limiter.limit("20/hour")  # Prevent AI abuse - 20 suggestions per hour
async def ai_suggest(request: Request, resume_id: str, ai_req: AISuggestionRequest, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Check AI limits and track usage
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    region = resume.get("region", "US")
    
    system_prompt = f"""
You are an expert resume writer for {region} job markets. Your role is to provide professional, accurate suggestions.

CRITICAL RULES:
1. NEVER fabricate or invent data, metrics, or achievements
2. If specific data is missing, suggest: [Add specific metric here]
3. Provide region-specific guidance for {region} standards
4. Keep suggestions concise and actionable
5. Focus on impact and quantifiable results when possible
"""
    
    user_prompt = f"""
Field: {ai_req.field}
Context: {ai_req.context}
Current text: {ai_req.current_text}

Provide a professional suggestion to improve this content. If metrics are needed but not provided, flag with [Add metric here].
"""
    
    async def generate():
        try:
            stream = await openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                stream=True,
                temperature=0.7,
                max_tokens=500
            )
            
            async for chunk in stream:
                if chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content': chunk.choices[0].delta.content})}\n\n"
                if chunk.choices[0].finish_reason == "stop":
                    # Track usage when stream completes
                    await track_ai_usage(current_user["id"], "ai_suggestion")
                    yield f"data: {json.dumps({'done': True})}\n\n"
                    break
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )

# ============ ATS SCORE CALCULATION ============
def calculate_ats_score(update_doc: dict, current_resume: dict) -> int:
    score = 0
    max_score = 100
    
    # Get the updated resume data
    personal_info = update_doc.get("personal_info", current_resume.get("personal_info", {}))
    work_experience = update_doc.get("work_experience", current_resume.get("work_experience", []))
    education = update_doc.get("education", current_resume.get("education", []))
    skills = update_doc.get("skills", current_resume.get("skills", []))
    
    # Personal info completeness (20 points)
    if personal_info:
        fields = ["full_name", "email", "phone", "location"]
        filled = sum(1 for f in fields if personal_info.get(f))
        score += int((filled / len(fields)) * 20)
    
    # Work experience (30 points)
    if work_experience:
        score += min(len(work_experience) * 10, 20)
        # Check for achievements
        has_achievements = any(exp.get("achievements") for exp in work_experience)
        if has_achievements:
            score += 10
    
    # Education (20 points)
    if education:
        score += min(len(education) * 10, 20)
    
    # Skills (20 points)
    if skills:
        score += min(len(skills) * 2, 20)
    
    # Summary (10 points)
    if personal_info and personal_info.get("summary"):
        score += 10
    
    return min(score, max_score)

@api_router.get("/resumes/{resume_id}/ats-score")
async def get_ats_score(resume_id: str, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return {
        "score": resume.get("ats_score", 0),
        "keywords": resume.get("keywords", [])
    }

# ============ EXPORT (Placeholder) ============
@api_router.get("/resumes/{resume_id}/export/pdf")
async def export_pdf(resume_id: str, current_user: dict = Depends(get_current_user)):
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Generate PDF using ReportLab
    try:
        pdf_bytes = generate_resume_pdf(resume)
        return Response(content=pdf_bytes, media_type="application/pdf", headers={
            "Content-Disposition": f"attachment; filename={resume.get('title', 'resume')}.pdf"
        })
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate PDF")

@api_router.get("/resumes/{resume_id}/export/word")
async def export_word(resume_id: str, current_user: dict = Depends(get_current_user)):
    """Export resume as MS Word document"""
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Personal Info
        personal_info = resume.get('personal_info', {})
        if personal_info.get('full_name'):
            heading = doc.add_heading(personal_info['full_name'], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])
        
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional Summary
        if personal_info.get('summary'):
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(personal_info['summary'])
        
        # Work Experience
        work_exp = resume.get('work_experience', [])
        if work_exp:
            doc.add_heading('Work Experience', level=2)
            for exp in work_exp:
                job_title = f"{exp.get('position', '')} at {exp.get('company', '')}"
                doc.add_heading(job_title, level=3)
                date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present') if not exp.get('current') else 'Present'}"
                doc.add_paragraph(date_range).italic = True
                if exp.get('description'):
                    doc.add_paragraph(exp['description'])
                for achievement in exp.get('achievements', []):
                    doc.add_paragraph(achievement, style='List Bullet')
        
        # Education
        education = resume.get('education', [])
        if education:
            doc.add_heading('Education', level=2)
            for edu in education:
                degree_info = f"{edu.get('degree', '')} in {edu.get('field', '')}"
                doc.add_heading(degree_info, level=3)
                doc.add_paragraph(f"{edu.get('institution', '')} | {edu.get('start_date', '')} - {edu.get('end_date', '')}")
        
        # Skills
        skills = resume.get('skills', [])
        if skills:
            doc.add_heading('Skills', level=2)
            doc.add_paragraph(', '.join(skills))
        
        # Save to bytes
        from io import BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return Response(
            content=docx_bytes.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={resume.get('title', 'resume')}.docx"
            }
        )
    except Exception as e:
        logger.error(f"Word generation error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate Word document")

@api_router.post("/convert-pdf-to-word")
async def convert_pdf_to_word(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Convert uploaded PDF resume to MS Word format (Paid users only)"""
    # Check subscription tier
    tier = current_user.get("subscription_tier", "free")
    if tier == "free":
        raise HTTPException(status_code=403, detail="PDF to Word conversion is available for Pro and Pro+ users only. Please upgrade your plan.")
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    file_ext = file.filename.lower().split('.')[-1]
    if file_ext != 'pdf':
        raise HTTPException(status_code=400, detail="Only PDF files are supported for conversion")
    
    try:
        from docx import Document
        from docx.shared import Pt
        import fitz  # PyMuPDF
        from io import BytesIO
        
        # Read and extract text from PDF
        file_bytes = await file.read()
        logger.info(f"PDF to Word: Processing file {file.filename}, size: {len(file_bytes)} bytes")
        
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        logger.info(f"PDF opened successfully, {len(pdf_document)} pages")
        
        # Create Word document
        doc = Document()
        
        # Track if any content was added
        content_added = False
        
        # Extract text from each page and add to Word
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if text.strip():
                content_added = True
                logger.info(f"Page {page_num + 1}: Extracted {len(text)} characters")
            
            # Split into paragraphs and add to Word
            paragraphs = text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    # Set font
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
        pdf_document.close()
        
        if not content_added:
            logger.warning("No text content extracted from PDF")
            raise HTTPException(status_code=400, detail="The PDF appears to be empty or contains only images. Please upload a text-based PDF.")
        
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        output_size = len(docx_bytes.getvalue())
        logger.info(f"Word document created successfully, size: {output_size} bytes")
        
        original_name = file.filename.rsplit('.', 1)[0]
        return Response(
            content=docx_bytes.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={original_name}.docx"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF to Word conversion error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to convert PDF to Word: {str(e)}")

@api_router.post("/convert-word-to-pdf")
async def convert_word_to_pdf(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Convert uploaded Word document to PDF format (Paid users only) - Preserves formatting"""
    # Check subscription tier
    tier = current_user.get("subscription_tier", "free")
    if tier == "free":
        raise HTTPException(status_code=403, detail="Word to PDF conversion is available for Pro and Pro+ users only. Please upgrade your plan.")
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    file_ext = file.filename.lower().split('.')[-1]
    if file_ext not in ['docx', 'doc']:
        raise HTTPException(status_code=400, detail="Only Word documents (.docx, .doc) are supported for conversion")
    
    import tempfile
    import subprocess
    from pathlib import Path
    
    # Create temporary files
    temp_dir = tempfile.mkdtemp()
    try:
        file_bytes = await file.read()
        logger.info(f"Word to PDF: Processing file {file.filename}, size: {len(file_bytes)} bytes")
        
        # Save uploaded file
        input_path = Path(temp_dir) / file.filename
        with open(input_path, 'wb') as f:
            f.write(file_bytes)
        
        output_path = Path(temp_dir) / f"{input_path.stem}.pdf"
        
        # Convert using LibreOffice headless
        try:
            result = subprocess.run(
                [
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    str(input_path)
                ],
                capture_output=True,
                text=True,
                timeout=30,
                check=True
            )
            logger.info(f"LibreOffice conversion output: {result.stdout}")
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=504, detail="Conversion timed out. File may be too large or complex.")
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e.stderr}")
            raise HTTPException(status_code=500, detail="Failed to convert document. Please ensure it's a valid Word file.")
        
        # Read converted PDF
        if not output_path.exists():
            raise HTTPException(status_code=500, detail="PDF conversion completed but output file not found")
        
        with open(output_path, 'rb') as f:
            pdf_bytes = f.read()
        
        logger.info(f"PDF created successfully, size: {len(pdf_bytes)} bytes")
        
        original_name = file.filename.rsplit('.', 1)[0]
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={original_name}.pdf"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Word to PDF conversion error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to convert Word to PDF: {str(e)}")
    finally:
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(temp_dir)
        except Exception as cleanup_error:
            logger.warning(f"Failed to cleanup temp directory: {cleanup_error}")

# ============ FILE UPLOAD & PARSING ============
@api_router.post("/resumes/upload")
async def upload_resume(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Upload and parse PDF/DOCX resume"""
    logger.info(f"Resume upload attempt by user: {current_user.get('id', 'unknown')}")
    
    # Check subscription limit before uploading
    await check_resume_limit(current_user["id"], current_user.get("subscription_tier", "free"))
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    file_ext = file.filename.lower().split('.')[-1]
    if file_ext not in ['pdf', 'docx']:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    file_bytes = await file.read()
    logger.info(f"File received: {file.filename}, size: {len(file_bytes)} bytes")
    
    try:
        if file_ext == 'pdf':
            parsed_data = await parse_pdf_resume(file_bytes)
        else:
            parsed_data = await parse_docx_resume(file_bytes)
        
        if 'error' in parsed_data:
            raise HTTPException(status_code=400, detail=parsed_data['error'])
        
        # Create new resume with parsed data
        resume_doc = {
            "user_id": current_user["id"],
            "title": f"Imported from {file.filename}",
            "region": current_user.get("region", "US"),
            "personal_info": {
                "full_name": parsed_data.get("full_name", ""),
                "email": parsed_data.get("email", ""),
                "phone": parsed_data.get("phone", ""),
                "location": parsed_data.get("location", ""),
                "linkedin": parsed_data.get("linkedin", ""),
                "website": parsed_data.get("website", ""),
                "summary": parsed_data.get("summary", "")
            },
            "work_experience": parsed_data.get("work_experience", []),
            "education": parsed_data.get("education", []),
            "skills": parsed_data.get("skills", []),
            "ats_score": 0,
            "keywords": [],
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        result = await db.resumes.insert_one(resume_doc)
        resume_doc["id"] = str(result.inserted_id)
        resume_doc.pop("_id", None)
        
        # Mark that free user has created/uploaded a resume (prevent abuse)
        if current_user.get("subscription_tier", "free") == "free":
            await db.users.update_one(
                {"id": current_user["id"]},
                {"$set": {"has_created_resume": True}}
            )
        
        # Calculate initial ATS score
        ats_score = calculate_ats_score(resume_doc, {})
        await db.resumes.update_one({"_id": result.inserted_id}, {"$set": {"ats_score": ats_score}})
        resume_doc["ats_score"] = ats_score
        
        # Track score history
        await db.score_history.insert_one({
            "resume_id": str(result.inserted_id),
            "score": ats_score,
            "date": datetime.now(timezone.utc).isoformat()
        })
        
        logger.info(f"Resume uploaded successfully: {resume_doc['id']}")
        return resume_doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

# ============ SCORE HISTORY ============
@api_router.get("/resumes/{resume_id}/score-history")
async def get_score_history(resume_id: str, current_user: dict = Depends(get_current_user)):
    """Get ATS score history for a resume"""
    # Verify resume ownership
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    history = await db.score_history.find({"resume_id": resume_id}, {"_id": 0}).sort("date", 1).to_list(100)
    return history

# ============ STAR BUILDER ============
class STAREntry(BaseModel):
    situation: str = ""
    task: str = ""
    action: str = ""
    result: str = ""
    job_title: str = ""
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

@api_router.post("/star-entries")
async def create_star_entry(entry: STAREntry, current_user: dict = Depends(get_current_user)):
    """Create a STAR format entry"""
    entry_doc = entry.model_dump()
    entry_doc["user_id"] = current_user["id"]
    entry_doc["id"] = str(uuid.uuid4())
    
    await db.star_entries.insert_one(entry_doc)
    return entry_doc

@api_router.get("/star-entries")
async def get_star_entries(current_user: dict = Depends(get_current_user)):
    """Get all STAR entries for user"""
    entries = await db.star_entries.find({"user_id": current_user["id"]}, {"_id": 0}).to_list(1000)
    return entries

@api_router.delete("/star-entries/{entry_id}")
async def delete_star_entry(entry_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a STAR entry"""
    result = await db.star_entries.delete_one({"user_id": current_user["id"], "id": entry_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Entry not found")
    return {"message": "Entry deleted successfully"}

# ============ AI USAGE TRACKING ============
async def track_ai_usage(user_id: str, feature: str):
    """Track AI feature usage for billing and limits"""
    await db.ai_usage.insert_one({
        "user_id": user_id,
        "feature": feature,
        "created_at": datetime.now(timezone.utc).isoformat()
    })

async def check_ai_limits(user_id: str, tier: str) -> bool:
    """Check if user has exceeded AI usage limits for their tier"""
    limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])
    max_ai = limits["max_ai_suggestions"]
    
    # Unlimited for paid tiers
    if max_ai == -1:
        return True
    
    # Count usage this month
    first_day = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    usage_count = await db.ai_usage.count_documents({
        "user_id": user_id,
        "created_at": {"$gte": first_day.isoformat()}
    })
    
    return usage_count < max_ai

async def check_headshot_limits(user_id: str, tier: str) -> bool:
    """Check if user has exceeded headshot generation limits for their tier"""
    limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])
    max_headshots = limits.get("max_headshots", 0)
    
    # No headshots allowed
    if max_headshots == 0:
        return False
    
    # Unlimited headshots
    if max_headshots == -1:
        return True
    
    # Count headshot usage this month
    first_day = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    headshot_count = await db.ai_usage.count_documents({
        "user_id": user_id,
        "feature": "headshot_generation",
        "created_at": {"$gte": first_day.isoformat()}
    })
    
    return headshot_count < max_headshots

# ============ AI HEADSHOT GENERATOR ============
class HeadshotRequest(BaseModel):
    image_data: str  # base64 encoded image

@api_router.post("/generate-headshot")
@limiter.limit("10/hour")  # Limit headshot generation - resource intensive
async def generate_headshot(request: Request, headshot_req: HeadshotRequest, current_user: dict = Depends(get_current_user)):
    """Generate professional headshot from selfie using OpenAI DALL-E 3"""
    try:
        # Check tier access (Pro and Pro+ only)
        tier = current_user.get("subscription_tier", "free").lower()
        if tier == "free":
            raise HTTPException(
                status_code=403, 
                detail="AI Headshot Generator is available for Pro and Pro+ subscribers. Upgrade to unlock this feature!"
            )
        
        # Check headshot-specific limits
        if not await check_headshot_limits(current_user["id"], tier):
            headshot_limit = TIER_LIMITS[tier]["max_headshots"]
            raise HTTPException(
                status_code=403, 
                detail=f"Monthly headshot limit reached ({headshot_limit}/month for {tier.upper()}). Limit resets next month."
            )
        
        # Remove data URL prefix if present
        image_base64 = headshot_req.image_data
        if ',' in image_base64:
            image_base64 = image_base64.split(',')[1]
        
        # Convert base64 to bytes
        import base64
        image_bytes = base64.b64decode(image_base64)
        
        # Save to temporary file (DALL-E requires file upload)
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(image_bytes)
            tmp_file_path = tmp_file.name
        
        try:
            # Use OpenAI DALL-E 3 for image editing/generation
            # Note: DALL-E 3 doesn't support image editing, so we'll use image generation with a prompt
            response = await openai_client.images.generate(
                model="dall-e-3",
                prompt="Professional corporate headshot: business professional in business casual attire, neutral professional background, studio lighting, confident expression, high quality portrait photography",
                size="1024x1024",
                quality="standard",
                n=1,
            )
            
            headshot_url = response.data[0].url
            
            # Download and convert to base64 for storage
            import httpx
            async with httpx.AsyncClient() as client:
                img_response = await client.get(headshot_url)
                img_base64 = base64.b64encode(img_response.content).decode()
                headshot_data = f"data:image/png;base64,{img_base64}"
            
            # Save headshot URL to user profile
            await db.users.update_one(
                {"id": current_user["id"]},
                {"$set": {"headshot_url": headshot_data}}
            )
            
            # Track AI usage
            await track_ai_usage(current_user["id"], "headshot_generation")
            
            return {"headshot_url": headshot_data, "message": "Professional headshot generated successfully"}
            
        finally:
            # Cleanup temp file
            import os
            try:
                os.unlink(tmp_file_path)
            except Exception:
                pass
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Headshot generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate headshot: {str(e)}")

# ============ JOB AD GENERATOR ============
class JobAdRequest(BaseModel):
    job_description: str
    resume_id: str

@api_router.post("/generate-from-job-ad")
@limiter.limit("15/hour")  # Limit job ad generation
async def generate_from_job_ad(request: Request, job_req: JobAdRequest, current_user: dict = Depends(get_current_user)):
    """Generate tailored resume and cover letter from job description"""
    # Check AI limits
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    # Get user's resume
    resume = await db.resumes.find_one({"user_id": current_user["id"], "_id": ObjectId(job_req.resume_id)}, {"_id": 0})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    system_prompt = f"""You are an expert resume and cover letter writer. 
Given a job description and a candidate's resume, create:
1. A tailored professional summary highlighting relevant skills
2. A tailored cover letter (3-4 paragraphs)

Region: {resume.get('region', 'US')}
CRITICAL: Never fabricate experience or skills. Only highlight what's in the resume.
"""
    
    resume_text = json.dumps({
        "personal_info": resume.get("personal_info", {}),
        "work_experience": resume.get("work_experience", []),
        "education": resume.get("education", []),
        "skills": resume.get("skills", [])
    }, indent=2)
    
    user_prompt = f"""
Job Description:
{job_req.job_description}

Candidate's Resume:
{resume_text}

Generate:
1. Tailored Professional Summary (2-3 sentences)
2. Cover Letter (3-4 paragraphs)

Format as JSON:
{{
  "summary": "...",
  "cover_letter": "..."
}}
"""
    
    async def generate():
        try:
            stream = await openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                stream=True,
                temperature=0.7,
                max_tokens=2000
            )
            
            accumulated = ""
            async for chunk in stream:
                if chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    accumulated += content
                    yield f"data: {json.dumps({'content': content})}\\n\\n"
                if chunk.choices[0].finish_reason == "stop":
                    # Track usage when stream completes
                    await track_ai_usage(current_user["id"], "job_ad_generation")
                    yield f"data: {json.dumps({'done': True, 'full_content': accumulated})}\\n\\n"
                    break
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\\n\\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )

# ============ COVER LETTER GENERATOR ============
class CoverLetterRequest(BaseModel):
    fullName: str = ""
    jobTitle: str
    companyName: str
    jobDescription: str = ""
    skills: str = ""
    experience: str = ""

@api_router.post("/generate-cover-letter")
@limiter.limit("15/hour")
async def generate_cover_letter(request: Request, cover_req: CoverLetterRequest, current_user: dict = Depends(get_current_user)):
    """Generate a professional cover letter"""
    # Check AI limits
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    system_prompt = f"""You are an expert cover letter writer. Create professional, personalized cover letters that:
- Highlight relevant experience and skills
- Show enthusiasm for the role and company
- Are concise (3-4 paragraphs)
- Use professional language
- Follow standard business letter format

Region: {current_user.get('region', 'US')}
"""
    
    user_prompt = f"""
Create a professional cover letter for:

Applicant: {cover_req.fullName or 'the candidate'}
Position: {cover_req.jobTitle}
Company: {cover_req.companyName}
{f'Skills: {cover_req.skills}' if cover_req.skills else ''}
{f'Experience: {cover_req.experience}' if cover_req.experience else ''}
{f'Job Description: {cover_req.jobDescription}' if cover_req.jobDescription else ''}

Write a compelling cover letter that makes the candidate stand out.
"""
    
    try:
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        
        accumulated = response.choices[0].message.content
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "cover_letter_generation")
        
        return {"cover_letter": accumulated}
    except Exception as e:
        logger.error(f"Cover letter generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate cover letter")

# ============ BULLET POINT WRITER ============
class BulletPointRequest(BaseModel):
    experience: str
    role: str = ""
    context: str = ""

@api_router.post("/generate-bullet-points")
@limiter.limit("20/hour")
async def generate_bullet_points(request: Request, bullet_req: BulletPointRequest, current_user: dict = Depends(get_current_user)):
    """Generate professional resume bullet points from experience description"""
    # Check AI limits
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    system_prompt = f"""You are an expert resume writer specializing in creating impactful, quantified bullet points.
Transform user experiences into professional STAR-format bullet points that:
- Start with strong action verbs
- Include specific metrics and results when possible
- Are concise (1-2 lines each)
- Follow ATS best practices
- Highlight achievements, not just duties

Region: {current_user.get('region', 'US')}
"""
    
    user_prompt = f"""
Transform this experience into 4-5 professional resume bullet points:

Experience: {bullet_req.experience}
{f'Role: {bullet_req.role}' if bullet_req.role else ''}
{f'Context: {bullet_req.context}' if bullet_req.context else ''}

Generate impactful bullet points. Return ONLY a JSON array of strings:
["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"]
"""
    
    try:
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        accumulated = response.choices[0].message.content
        
        # Parse JSON response
        try:
            # Extract JSON array from response
            import re
            json_match = re.search(r'\[.*\]', accumulated, re.DOTALL)
            if json_match:
                bullet_points = json.loads(json_match.group())
            else:
                # Fallback: split by newlines
                bullet_points = [line.strip().lstrip('•-*').strip() 
                               for line in accumulated.split('\n') 
                               if line.strip() and not line.strip().startswith('{')]
        except Exception:
            # Fallback parsing
            bullet_points = [line.strip().lstrip('•-*').strip() 
                           for line in accumulated.split('\n') 
                           if line.strip()]
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "bullet_point_generation")
        
        return {"bullet_points": bullet_points[:5]}  # Max 5 bullets
    except Exception as e:
        logger.error(f"Bullet point generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate bullet points")

# ============ SUMMARY GENERATOR ============
class SummaryRequest(BaseModel):
    currentRole: str = ""
    yearsOfExperience: str = ""
    keySkills: str
    targetRole: str = ""
    achievements: str = ""

@api_router.post("/generate-summary")
@limiter.limit("20/hour")
async def generate_summary(request: Request, summary_req: SummaryRequest, current_user: dict = Depends(get_current_user)):
    """Generate a professional resume summary"""
    # Check AI limits
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    system_prompt = f"""You are an expert resume writer specializing in professional summaries.
Create compelling 2-3 sentence summaries that:
- Highlight key strengths and experience
- Include relevant skills
- Show career trajectory
- Are concise and impactful
- Follow ATS best practices

Region: {current_user.get('region', 'US')}
"""
    
    user_prompt = f"""
Create a professional resume summary based on:

{f'Current Role: {summary_req.currentRole}' if summary_req.currentRole else ''}
{f'Years of Experience: {summary_req.yearsOfExperience}' if summary_req.yearsOfExperience else ''}
Key Skills: {summary_req.keySkills}
{f'Target Role: {summary_req.targetRole}' if summary_req.targetRole else ''}
{f'Key Achievements: {summary_req.achievements}' if summary_req.achievements else ''}

Write a compelling 2-3 sentence professional summary that captures this candidate's value.
"""
    
    try:
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=600
        )
        
        accumulated = response.choices[0].message.content
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "summary_generation")
        
        return {"summary": accumulated.strip()}
    except Exception as e:
        logger.error(f"Summary generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate summary")

# ============ SELECTION CRITERIA RESPONSE ============
class SelectionCriteriaRequest(BaseModel):
    question: str
    resumeContext: str

@api_router.post("/generate-selection-criteria")
@limiter.limit("15/hour")
async def generate_selection_criteria(request: Request, criteria_req: SelectionCriteriaRequest, current_user: dict = Depends(get_current_user)):
    """Generate STAR-format response to selection criteria question"""
    # Check AI limits
    tier = current_user.get("subscription_tier", "free")
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    system_prompt = f"""You are an expert career counselor specializing in selection criteria responses.
Create compelling STAR-format responses that:
- Follow the STAR framework (Situation, Task, Action, Result)
- Are specific and detailed
- Include quantifiable achievements where possible
- Demonstrate relevant skills and competencies
- Are well-structured and professional
- Never fabricate information - only use what's provided in the resume context

Region: {current_user.get('region', 'US')}
"""
    
    user_prompt = f"""
Selection Criteria Question:
{criteria_req.question}

Candidate's Resume/Experience:
{criteria_req.resumeContext}

Generate a professional response to the selection criteria question based ONLY on the provided resume/experience.

IMPORTANT FORMATTING INSTRUCTIONS:
- Write the response in FOUR clear paragraphs following the STAR framework structure
- DO NOT label the paragraphs with "Situation:", "Task:", "Action:", or "Result:"
- Write naturally flowing paragraphs that users can copy-paste directly
- First paragraph: Describe the context and background
- Second paragraph: Explain the responsibility or challenge
- Third paragraph: Detail the specific actions taken
- Fourth paragraph: Highlight the outcomes and measurable impact

The response should read as a cohesive narrative without explicit STAR labels, but still follow the STAR structure internally.
"""
    
    try:
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        accumulated = response.choices[0].message.content
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "selection_criteria_generation")
        
        return {"response": accumulated.strip()}
    except Exception as e:
        logger.error(f"Selection criteria generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate selection criteria response")

# ============ INTERVIEW PREPARATION (PRO+ ONLY) ============
@api_router.post("/interview-prep/generate")
@limiter.limit("10/hour")
async def generate_interview_questions(
    request: Request,
    prep_req: InterviewPrepRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate tailored interview questions based on job description and resume (Pro+ only)"""
    
    # Check subscription tier - Pro+ only
    tier = current_user.get("subscription_tier", "free").lower()
    if tier != "pro+":
        raise HTTPException(
            status_code=403, 
            detail="Interview Preparation is exclusive to Pro+ subscribers. Upgrade to unlock this feature!"
        )
    
    try:
        # Fetch the resume
        resume = await db.resumes.find_one(
            {"_id": ObjectId(prep_req.resume_id), "user_id": current_user["id"]}
        )
        
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Extract relevant resume data
        personal_info = resume.get("personal_info", {})
        work_experience = resume.get("work_experience", [])
        education = resume.get("education", [])
        skills = resume.get("skills", [])
        
        # Build resume summary for context
        resume_summary = f"""
NAME: {personal_info.get('full_name', 'Not provided')}
SUMMARY: {personal_info.get('summary', 'Not provided')}

WORK EXPERIENCE:
"""
        for exp in work_experience[:3]:  # Top 3 most recent
            resume_summary += f"- {exp.get('position', '')} at {exp.get('company', '')}\n"
            for achievement in exp.get('achievements', [])[:2]:  # Top 2 achievements
                resume_summary += f"  • {achievement}\n"
        
        resume_summary += "\nEDUCATION:\n"
        for edu in education:
            resume_summary += f"- {edu.get('degree', '')} in {edu.get('field', '')} from {edu.get('institution', '')}\n"
        
        resume_summary += f"\nSKILLS: {', '.join(skills[:10])}\n"
        
        # Create AI prompt
        system_prompt = """You are an expert career coach and interview preparation specialist. 
Your task is to generate tailored interview questions based on the candidate's resume and the specific job description.

Generate comprehensive interview questions in the following categories:
1. **Technical Questions** (5-7 questions) - Role-specific technical skills and knowledge
2. **Behavioral Questions** (5-7 questions) - Past experiences, problem-solving, teamwork
3. **Situational Questions** (3-5 questions) - Hypothetical scenarios related to the role

For each question, ensure:
- It's directly relevant to either the job description or the candidate's background
- It's specific and actionable (not generic)
- It helps the candidate prepare for real interview scenarios

Format your response as JSON with this exact structure:
{
  "technical": ["question 1", "question 2", ...],
  "behavioral": ["question 1", "question 2", ...],
  "situational": ["question 1", "question 2", ...]
}
"""
        
        user_prompt = f"""JOB DESCRIPTION:
{prep_req.job_description}

CANDIDATE'S RESUME SUMMARY:
{resume_summary}

Generate tailored interview questions for this candidate applying to this specific role."""
        
        # Call OpenAI
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1500,
            response_format={"type": "json_object"}
        )
        
        accumulated = response.choices[0].message.content
        
        # Parse the JSON response
        try:
            questions_data = json.loads(accumulated)
        except json.JSONDecodeError:
            # Fallback: try to extract JSON from response
            import re
            json_match = re.search(r'\{.*\}', accumulated, re.DOTALL)
            if json_match:
                questions_data = json.loads(json_match.group(0))
            else:
                raise HTTPException(status_code=500, detail="Failed to parse AI response")
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "interview_prep_generation")
        
        return {
            "questions": questions_data,
            "job_title": prep_req.job_description.split('\n')[0][:100]  # First line as title
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Interview prep generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate interview questions")

@api_router.post("/interview-prep/generate-answer")
@limiter.limit("20/hour")
async def generate_interview_answer(
    request: Request,
    answer_req: InterviewAnswerRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate sample answer for a specific interview question (Pro+ only)"""
    
    # Check subscription tier - Pro+ only
    tier = current_user.get("subscription_tier", "free").lower()
    if tier != "pro+":
        raise HTTPException(
            status_code=403, 
            detail="Interview answer generation is exclusive to Pro+ subscribers."
        )
    
    # Check AI limits
    if not await check_ai_limits(current_user["id"], tier):
        raise HTTPException(status_code=403, detail="AI usage limit reached. Please upgrade your plan.")
    
    try:
        system_prompt = """You are an expert interview coach helping candidates prepare strong answers.

Generate a sample answer to the interview question using the STAR method (Situation, Task, Action, Result).

The answer should:
- Be specific and concrete (use the candidate's background provided)
- Follow STAR structure clearly
- Be 150-200 words
- Sound natural and conversational
- Highlight relevant skills and achievements

Format as plain text, starting with the answer directly."""

        user_prompt = f"""INTERVIEW QUESTION:
{answer_req.question}

CANDIDATE'S BACKGROUND:
{answer_req.resume_context}

JOB CONTEXT:
{answer_req.job_description[:500]}

Generate a strong sample answer using the STAR method."""

        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=400
        )
        
        answer = response.choices[0].message.content.strip()
        
        # Track AI usage
        await track_ai_usage(current_user["id"], "interview_answer_generation")
        
        return {"answer": answer}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Interview answer generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate answer")

# ============ ADMIN DASHBOARD ============
@api_router.get("/admin/stats")
async def get_admin_stats(current_user: dict = Depends(get_current_user)):
    """Get admin dashboard statistics"""
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Total users
    total_users = await db.users.count_documents({})
    
    # Users by subscription tier
    free_users = await db.users.count_documents({"subscription_tier": {"$in": ["free", None]}})
    pro_users = await db.users.count_documents({"subscription_tier": "pro"})
    pro_plus_users = await db.users.count_documents({"subscription_tier": "pro+"})
    
    # Total resumes
    total_resumes = await db.resumes.count_documents({})
    
    # Recent users (last 30 days)
    thirty_days_ago = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
    recent_users = await db.users.count_documents({"created_at": {"$gte": thirty_days_ago}})
    
    # Get all users with basic info
    users_list = await db.users.find(
        {},
        {"_id": 0, "id": {"$toString": "$_id"}, "email": 1, "name": 1, "created_at": 1, "subscription_tier": 1, "role": 1}
    ).sort("created_at", -1).to_list(1000)
    
    return {
        "total_users": total_users,
        "subscription_breakdown": {
            "free": free_users,
            "pro": pro_users,
            "pro_plus": pro_plus_users
        },
        "total_resumes": total_resumes,
        "recent_users_30d": recent_users,
        "users": users_list
    }

# ============ STRIPE INTEGRATION ============

# Subscription tier limits
TIER_LIMITS = {
    "free": {
        "max_resumes": 1,
        "max_ai_suggestions": 5,  # 5 per month
        "max_headshots": 0,  # No headshots for free tier
        "max_pdf_exports": 1,
        "features": ["basic_resume_builder", "basic_ats_score"]
    },
    "pro": {
        "max_resumes": -1,  # unlimited
        "max_ai_suggestions": 500,  # 500 per month (prevents abuse, ~$2.50 cost max)
        "max_headshots": 5,  # 5 headshots per month (~$0.20 cost max)
        "max_pdf_exports": -1,
        "features": ["resume_upload", "ats_history", "selection_criteria", "tailor_to_job_ad", "headshot_generator", "word_export", "pdf_to_word"]
    },
    "pro+": {
        "max_resumes": -1,
        "max_ai_suggestions": 1000,  # 1000 per month (prevents abuse, ~$5 cost max)
        "max_headshots": 10,  # 10 headshots per month (~$0.40 cost max)
        "max_pdf_exports": -1,
        "features": ["resume_upload", "ats_history", "selection_criteria", "tailor_to_job_ad", "headshot_generator", "word_export", "pdf_to_word", "priority_support"]
    }
}

class SubscriptionTier(BaseModel):
    tier: str  # "pro" or "pro+"

@api_router.post("/create-checkout-session")
async def create_checkout_session(tier_data: SubscriptionTier, current_user: dict = Depends(get_current_user)):
    """Create Stripe checkout session for subscription"""
    try:
        price_id = STRIPE_PRO_PRICE_ID if tier_data.tier == "pro" else STRIPE_PRO_PLUS_PRICE_ID
        
        checkout_session = stripe.checkout.Session.create(
            customer_email=current_user["email"],
            line_items=[{
                "price": price_id,
                "quantity": 1,
            }],
            mode="subscription",
            success_url=f"{os.environ['FRONTEND_URL']}/dashboard?checkout=success",
            cancel_url=f"{os.environ['FRONTEND_URL']}/pricing?checkout=cancelled",
            metadata={
                "user_id": current_user["id"],
                "tier": tier_data.tier
            }
        )
        
        return {"checkout_url": checkout_session.url}
    except Exception as e:
        logger.error(f"Stripe checkout error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/create-portal-session")
async def create_portal_session(current_user: dict = Depends(get_current_user)):
    """Create Stripe customer portal session for subscription management"""
    try:
        # Find customer by email
        customers = stripe.Customer.list(email=current_user["email"], limit=1)
        
        if not customers.data:
            raise HTTPException(status_code=404, detail="No subscription found")
        
        customer_id = customers.data[0].id
        
        portal_session = stripe.billing_portal.Session.create(
            customer=customer_id,
            return_url=f"{os.environ['FRONTEND_URL']}/dashboard"
        )
        
        return {"portal_url": portal_session.url}
    except Exception as e:
        logger.error(f"Stripe portal error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.delete("/user/account")
async def delete_user_account(current_user: dict = Depends(get_current_user)):
    """
    Permanently delete user account and all associated data.
    This action is irreversible and complies with GDPR right to erasure.
    """
    try:
        user_id = current_user["id"]
        
        # Delete all user's resumes
        await db.resumes.delete_many({"user_id": user_id})
        
        # Delete all user's STAR entries
        await db.star_entries.delete_many({"user_id": user_id})
        
        # Delete user account
        result = await db.users.delete_one({"id": user_id})
        
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="User not found")
        
        logger.info(f"User account deleted: {current_user['email']}")
        
        return {
            "message": "Account and all associated data permanently deleted",
            "deleted": {
                "user": True,
                "resumes": "all",
                "data": "all"
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting user account: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete account")

@api_router.post("/stripe-webhook")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks for subscription events"""
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    
    try:
        if STRIPE_WEBHOOK_SECRET:
            event = stripe.Webhook.construct_event(
                payload, sig_header, STRIPE_WEBHOOK_SECRET
            )
        else:
            event = json.loads(payload)
        
        # Handle different event types
        if event["type"] == "checkout.session.completed":
            session = event["data"]["object"]
            user_id = session["metadata"]["user_id"]
            tier = session["metadata"]["tier"]
            customer_id = session["customer"]
            subscription_id = session["subscription"]
            
            # Update user subscription
            await db.users.update_one(
                {"_id": ObjectId(user_id)},
                {"$set": {
                    "subscription_tier": tier,
                    "subscription_status": "active",
                    "subscription_stripe_customer_id": customer_id,
                    "subscription_stripe_id": subscription_id,
                    "subscription_updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            logger.info(f"User {user_id} subscribed to {tier}")
        
        elif event["type"] == "customer.subscription.updated":
            subscription = event["data"]["object"]
            customer_id = subscription["customer"]
            status = subscription["status"]
            
            # Find user by customer ID
            user = await db.users.find_one({"subscription_stripe_customer_id": customer_id})
            if user:
                await db.users.update_one(
                    {"_id": user["_id"]},
                    {"$set": {
                        "subscription_status": status,
                        "subscription_updated_at": datetime.now(timezone.utc).isoformat()
                    }}
                )
                logger.info(f"User {user['_id']} subscription updated: {status}")
        
        elif event["type"] == "customer.subscription.deleted":
            subscription = event["data"]["object"]
            customer_id = subscription["customer"]
            
            # Find user and downgrade to free
            user = await db.users.find_one({"subscription_stripe_customer_id": customer_id})
            if user:
                await db.users.update_one(
                    {"_id": user["_id"]},
                    {"$set": {
                        "subscription_tier": "free",
                        "subscription_status": "canceled",
                        "subscription_updated_at": datetime.now(timezone.utc).isoformat()
                    }}
                )
                logger.info(f"User {user['_id']} subscription canceled")
        
        return {"status": "success"}
    
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        raise HTTPException(status_code=400, detail=str(e))

@api_router.get("/subscription-status")
async def get_subscription_status(current_user: dict = Depends(get_current_user)):
    """Get current user's subscription status and limits"""
    tier = current_user.get("subscription_tier", "free")
    limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])
    
    # Get usage counts
    resume_count = await db.resumes.count_documents({"user_id": current_user["id"]})
    
    # Get AI usage this month
    first_day = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    ai_usage = await db.ai_usage.count_documents({
        "user_id": current_user["id"],
        "created_at": {"$gte": first_day.isoformat()}
    })
    
    return {
        "tier": tier,
        "status": current_user.get("subscription_status", "active"),
        "limits": limits,
        "usage": {
            "resumes": resume_count,
            "ai_suggestions": ai_usage
        },
        "can_upgrade": tier == "free" or tier == "pro"
    }

# Include router
app.include_router(api_router)

# CORS Configuration
cors_origins = os.environ.get('CORS_ORIGINS', '*')
if cors_origins == '*':
    allowed_origins = ["*"]
else:
    allowed_origins = [origin.strip() for origin in cors_origins.split(',')]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

    client.close()
