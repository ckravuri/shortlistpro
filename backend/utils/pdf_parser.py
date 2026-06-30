import fitz
import io
from typing import Dict, List
import re
import json
import os

async def parse_pdf_resume(pdf_bytes: bytes) -> Dict:
    """
    Parse resume from PDF bytes using PyMuPDF
    Returns structured data: name, email, phone, skills, experience, education
    """
    text = ""
    
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in pdf_document:
            text += page.get_text()
        pdf_document.close()
    except Exception as e:
        return {"error": f"Failed to parse PDF: {str(e)}"}
    
    return await extract_resume_data_ai(text)

async def parse_docx_resume(docx_bytes: bytes) -> Dict:
    """
    Parse resume from DOCX bytes - extracts from paragraphs, headers, and tables
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        
        # Extract text from paragraphs
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Extract text from headers (many resumes have contact info in headers)
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    text = para.text + "\n" + text  # Prepend header text
        
        # Extract text from tables (some resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += "\n" + row_text
        
        return await extract_resume_data_ai(text)
    except Exception as e:
        return {"error": f"Failed to parse DOCX: {str(e)}"}

async def extract_resume_data_ai(text: str) -> Dict:
    """
    Async implementation of AI resume parsing
    """
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # Initialize LLM chat with Emergent LLM Key
        api_key = os.environ.get("EMERGENT_LLM_KEY")
        if not api_key:
            print("EMERGENT_LLM_KEY not found, falling back to basic extraction")
            return extract_resume_data_fallback(text)
        
        chat = LlmChat(
            api_key=api_key,
            session_id="resume_parser",
            system_message="You are a professional resume parser that extracts structured data and returns ONLY valid JSON."
        ).with_model("openai", "gpt-5.2")
        
        prompt = f"""Extract structured information from the following resume text and return ONLY a valid JSON object with NO markdown formatting, NO code blocks, NO extra text.

Resume text:
{text[:20000]}

Return a JSON object with this EXACT structure:
{{
  "full_name": "string (CRITICAL: Extract the person's full name - look at the very top of the resume, in headers, or near email/phone)",
  "email": "string (CRITICAL: Extract email address - look for @ symbol, usually at top of resume)",
  "phone": "string (CRITICAL: Extract phone number - look for digits, +, parentheses)",
  "location": "string (Extract city, state/country - look near name/contact info)",
  "linkedin": "string (URL or username)",
  "website": "string",
  "summary": "string (professional summary/objective/profile)",
  "skills": ["skill1", "skill2", "skill3"],
  "work_experience": [
    {{
      "id": "1",
      "position": "job title",
      "company": "company name",
      "location": "city, state",
      "start_date": "Month YYYY",
      "end_date": "Month YYYY or Present",
      "current": true/false,
      "description": "brief description",
      "achievements": ["achievement 1", "achievement 2"]
    }}
  ],
  "education": [
    {{
      "id": "1",
      "degree": "full degree name",
      "institution": "university/college name",
      "field": "field of study",
      "location": "city, state",
      "start_date": "YYYY",
      "end_date": "YYYY",
      "gpa": "X.X"
    }}
  ]
}}

CRITICAL RULES:
1. Return ONLY the JSON object, no markdown, no code blocks, no explanations
2. **IMPORTANT**: Look at the VERY BEGINNING of the resume text for full_name, email, phone, location
3. The name is usually the FIRST thing in a resume, often in larger font or at the top
4. Email contains @ symbol, phone contains numbers
5. If a field is not found, use empty string "" or empty array []
6. Extract ALL work experiences and education entries
7. For dates, normalize to "Month YYYY" format (e.g., "January 2020")
8. For current positions, set end_date to "Present" and current to true
9. Keep achievements concise (under 100 chars each)
10. If work experience has bullet points, put them in achievements array
11. Extract top 20 most relevant skills"""

        user_message = UserMessage(text=prompt)
        
        # Use non-streaming for parsing
        response = await chat.send_message(user_message)
        
        # Response is already a string
        result_text = response.strip() if isinstance(response, str) else response.content.strip()
        
        # Remove markdown code blocks if present
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
            result_text = result_text.strip()
        
        # Parse JSON
        parsed_data = json.loads(result_text)
        
        # Add raw_text
        parsed_data["raw_text"] = text
        
        return parsed_data
        
    except Exception as e:
        # Fallback to basic extraction if AI fails
        print(f"AI parsing failed: {e}, falling back to basic extraction")
        import traceback
        traceback.print_exc()
        return extract_resume_data_fallback(text)

def extract_resume_data_fallback(text: str) -> Dict:
    """
    Fallback extraction using basic patterns when AI is unavailable
    """
    data = {
        "full_name": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "website": "",
        "summary": "",
        "skills": [],
        "work_experience": [],
        "education": [],
        "raw_text": text
    }
    
    # Extract email
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        data["email"] = email_match.group()
    
    # Extract phone
    phone_match = re.search(r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]', text)
    if phone_match:
        phone_candidate = phone_match.group().strip()
        if not re.match(r'^\d{4}$', phone_candidate.replace(' ', '').replace('-', '')):
            data["phone"] = phone_candidate
    
    # Extract name (first substantial line)
    lines = text.split('\n')
    for line in lines[:10]:
        line = line.strip()
        if len(line) > 3 and len(line) < 60 and '@' not in line:
            data["full_name"] = line
            break
    
    # Extract skills - look for common section headers
    skills_pattern = r'(?:SKILLS?|TECHNICAL|COMPETENCIES)[:\s]*\n(.*?)(?=\n[A-Z]{4,}|\n\n|$)'
    skills_match = re.search(skills_pattern, text, re.IGNORECASE | re.DOTALL)
    if skills_match:
        skills_text = skills_match.group(1)[:500]
        skills = re.split(r'[,;•\|●◆▪\n]', skills_text)
        data["skills"] = [s.strip() for s in skills if s.strip() and len(s.strip()) > 2][:20]
    
    return data
